"""Generate shading.docx fixture for shading parser tests.

Run:
    python tests/fixtures/generate_shading.py
"""

from __future__ import annotations

from pathlib import Path
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

FIXTURES_DIR = Path(__file__).parent
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _shd(fill: str) -> etree._Element:
    el = etree.Element(qn("w:shd"))
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), fill)
    return el


def make_shading() -> None:
    doc = Document()
    # Remove default empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Paragraph with blue shading (4472C4)
    p1 = doc.add_paragraph("Blue shaded paragraph")
    ppr1 = p1._element.get_or_add_pPr()
    ppr1.append(_shd("4472C4"))

    # Paragraph with no shading
    doc.add_paragraph("No shading paragraph")

    # Table: first cell orange (ED7D31), second cell none
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    tc0 = table.rows[0].cells[0]._tc
    tc0pr = tc0.get_or_add_tcPr()
    tc0pr.append(_shd("ED7D31"))
    table.rows[0].cells[0].text = "Orange cell"
    table.rows[0].cells[1].text = "Plain cell"

    path = FIXTURES_DIR / "shading.docx"
    doc.save(str(path))
    print(f"  created  {path.name}")


if __name__ == "__main__":
    make_shading()
