"""
Generate binary DOCX fixture files for docwow parser tests.

Run once to create or recreate all fixtures:
    python tests/fixtures/generate_fixtures.py

Requires: python-docx, Pillow (Pillow is a core dev dep; python-docx is dev-only)

Each fixture targets specific parser code paths:
    empty.docx          — empty document, default page geometry
    paragraphs.docx     — Normal, Heading1–3, alignments
    formatting.docx     — bold, italic, underline, strike, font name/size/color,
                          superscript, subscript
    table_simple.docx   — plain 3×3 table
    table_merged.docx   — horizontal and vertical cell merges
    list_bullet.docx    — single-level bullet list
    list_numbered.docx  — single-level numbered list
    list_nested.docx    — 3-level nested bullet list
    image_inline.docx   — inline image inside a paragraph
    mixed.docx          — all element types combined
"""

import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage

FIXTURES_DIR = Path(__file__).parent


def _save(doc: Document, name: str) -> None:
    path = FIXTURES_DIR / name
    doc.save(str(path))
    print(f"  created  {path.name}")


def _red_png(width: int = 200, height: int = 100) -> io.BytesIO:
    buf = io.BytesIO()
    PILImage.new("RGB", (width, height), color=(255, 0, 0)).save(buf, format="PNG")
    buf.seek(0)
    return buf


def _green_png(width: int = 100, height: int = 50) -> io.BytesIO:
    buf = io.BytesIO()
    PILImage.new("RGB", (width, height), color=(0, 128, 0)).save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------

def make_empty() -> None:
    """Minimal document — exercises default page size / margin parsing."""
    doc = Document()
    # Remove any placeholder paragraphs python-docx adds by default
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    _save(doc, "empty.docx")


def make_paragraphs() -> None:
    """Normal paragraphs plus Heading 1–3 in all four alignments."""
    doc = Document()
    doc.add_heading("Heading 1", level=1)
    doc.add_heading("Heading 2", level=2)
    doc.add_heading("Heading 3", level=3)
    doc.add_paragraph("Normal paragraph with default alignment.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Left aligned paragraph.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Centered paragraph.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Right aligned paragraph.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Justified paragraph with enough text to demonstrate "
        "how the justify alignment works across the full line width."
    )
    _save(doc, "paragraphs.docx")


def make_formatting() -> None:
    """Inline character formatting: bold, italic, underline, strike,
    font name, font size, color, superscript, subscript."""
    doc = Document()

    # All toggles in one paragraph
    p = doc.add_paragraph()
    p.add_run("Normal ")
    p.add_run("Bold").bold = True
    p.add_run(" ")
    p.add_run("Italic").italic = True
    p.add_run(" ")
    r = p.add_run("Underline")
    r.underline = True
    p.add_run(" ")
    r = p.add_run("Strike")
    r.font.strike = True

    # Font size
    p2 = doc.add_paragraph()
    r = p2.add_run("Size 8pt")
    r.font.size = Pt(8)
    p2.add_run("  ")
    r = p2.add_run("Size 24pt")
    r.font.size = Pt(24)

    # Color
    p3 = doc.add_paragraph()
    r = p3.add_run("Red text")
    r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p3.add_run("  ")
    r = p3.add_run("Blue text")
    r.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    # Font name
    p4 = doc.add_paragraph()
    r = p4.add_run("Arial font")
    r.font.name = "Arial"
    p4.add_run("  ")
    r = p4.add_run("Courier New font")
    r.font.name = "Courier New"

    # Superscript / subscript
    p5 = doc.add_paragraph()
    p5.add_run("H")
    r = p5.add_run("2")
    r.font.subscript = True
    p5.add_run("O   E=mc")
    r = p5.add_run("2")
    r.font.superscript = True

    _save(doc, "formatting.docx")


def make_table_simple() -> None:
    """Plain 3×3 table — exercises basic cell / row / column parsing."""
    doc = Document()
    doc.add_paragraph("Paragraph before table.")
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"R{i + 1}C{j + 1}"
    doc.add_paragraph("Paragraph after table.")
    _save(doc, "table_simple.docx")


def make_table_merged() -> None:
    """Table with both horizontal (colspan) and vertical (rowspan) merges."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    # Fill every cell first so none are blank
    labels = [
        ["A", "B", "C"],
        ["D", "E", "F"],
        ["G", "H", "I"],
    ]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = labels[i][j]

    # Horizontal merge: row 0, col 0–1  →  "AB" spanning two columns
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "AB (colspan 2)"

    # Vertical merge: col 2, row 0–1  →  "CF" spanning two rows
    table.cell(0, 2).merge(table.cell(1, 2))
    table.cell(0, 2).text = "CF (rowspan 2)"

    _save(doc, "table_merged.docx")


def make_list_bullet() -> None:
    """Single-level unordered (bullet) list."""
    doc = Document()
    for i in range(1, 6):
        doc.add_paragraph(f"Bullet item {i}", style="List Bullet")
    _save(doc, "list_bullet.docx")


def make_list_numbered() -> None:
    """Single-level ordered (numbered) list."""
    doc = Document()
    for i in range(1, 6):
        doc.add_paragraph(f"Numbered item {i}", style="List Number")
    _save(doc, "list_numbered.docx")


def make_list_nested() -> None:
    """Three-level nested bullet list using built-in List Bullet 1/2/3 styles."""
    doc = Document()
    items = [
        (1, "Level 1 — Item A"),
        (2, "Level 2 — Item A.1"),
        (2, "Level 2 — Item A.2"),
        (3, "Level 3 — Item A.2.a"),
        (1, "Level 1 — Item B"),
        (2, "Level 2 — Item B.1"),
        (3, "Level 3 — Item B.1.a"),
        (3, "Level 3 — Item B.1.b"),
        (1, "Level 1 — Item C"),
    ]
    style_map = {
        1: "List Bullet",
        2: "List Bullet 2",
        3: "List Bullet 3",
    }
    for level, text in items:
        doc.add_paragraph(text, style=style_map[level])
    _save(doc, "list_nested.docx")


def make_image_inline() -> None:
    """Paragraph containing a single inline image."""
    doc = Document()
    doc.add_paragraph("Paragraph before the image.")
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(_red_png(), width=Inches(2))
    doc.add_paragraph("Paragraph after the image.")
    _save(doc, "image_inline.docx")


def make_mixed() -> None:
    """All element types in a single document — used for integration tests."""
    doc = Document()

    doc.add_heading("Mixed Document", level=1)
    doc.add_paragraph(
        "This document exercises all v0.1 element types: "
        "paragraphs, headings, inline formatting, tables, lists, and images."
    )

    doc.add_heading("Inline Formatting", level=2)
    p = doc.add_paragraph()
    p.add_run("Normal ")
    p.add_run("Bold").bold = True
    p.add_run(" ")
    p.add_run("Italic").italic = True
    p.add_run(" ")
    r = p.add_run("Colored")
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_heading("Table", level=2)
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = ["Name", "Role", "Status"]
    values = ["Alice", "Engineer", "Active"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    for j, v in enumerate(values):
        table.cell(1, j).text = v

    doc.add_heading("List", level=2)
    for item in ["Alpha", "Beta", "Gamma"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Image", level=2)
    p = doc.add_paragraph()
    p.add_run().add_picture(_green_png(), width=Inches(1))

    _save(doc, "mixed.docx")


# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating DOCX fixtures …")
    make_empty()
    make_paragraphs()
    make_formatting()
    make_table_simple()
    make_table_merged()
    make_list_bullet()
    make_list_numbered()
    make_list_nested()
    make_image_inline()
    make_mixed()
    print("Done.")
