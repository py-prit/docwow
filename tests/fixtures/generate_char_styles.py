"""Generate char_styles.docx fixture for character style parser tests.

Run:
    python tests/fixtures/generate_char_styles.py
"""

from __future__ import annotations

from pathlib import Path
from docx import Document

FIXTURES_DIR = Path(__file__).parent


def make_char_styles() -> None:
    doc = Document()
    # Remove default empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Paragraph with mixed character styles
    p = doc.add_paragraph()
    p.add_run("Plain text. ")
    r_strong = p.add_run("Strong text.")
    r_strong.style = doc.styles["Strong"]
    p.add_run(" ")
    r_em = p.add_run("Emphasized text.")
    r_em.style = doc.styles["Emphasis"]

    path = FIXTURES_DIR / "char_styles.docx"
    doc.save(str(path))
    print(f"  created  {path.name}")


if __name__ == "__main__":
    make_char_styles()
